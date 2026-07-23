# -*- coding: utf-8 -*-
"""
Hebrew CV + cover-letter rendering, by FILLING the user's own Hebrew Word CV
(config "hebrew_cv_template") rather than building a docx from scratch.

Why fill instead of generate: Word's RTL/bidi is finicky, and a from-scratch
python-docx file (any recipe we tried) garbles in real Word even when it looks
fine in the docx2pdf preview. The one thing that renders correctly in Word is
the user's own file's paragraph structure. So we keep the template's paragraphs (their
pPr carries the settings that make bidi work) and only swap the RUN TEXT inside
them. Verified against docx2pdf, which is faithful for files built this way.

Run-text rule (matches how the template splits runs):
  - split each line into per-script runs; Hebrew runs get <w:rtl/>, Latin runs
    don't; every run sets the complex-script font (Arial) + size.
  - digits are NEUTRAL (stay with the adjacent Hebrew run) so "14+"/"99%" render
    like the template. EXCEPTION: a digit glued to following Latin letters
    ("3PL") is merged into that Latin run so it does not reverse to "PL3".
  - a space at an LTR->RTL run boundary is moved onto the RTL run (Word trims it
    otherwise, gluing the words).

Public API:
    render_cv_from_template(cv_spec, out_path)
    render_cover_from_template(cover_spec, out_path)
"""
import os
import re
import copy
from docx import Document
from docx.oxml.ns import qn

_HEB = re.compile(r"[֐-׿]")
_TRAIL_DIGITS = re.compile(r"\d+$")
FONT = "Arial"

# No default template ships with the skill: the template is the USER'S own
# Hebrew Word CV (config "hebrew_cv_template"). Callers pass template_path.
TEMPLATE_PATH = None

# Role headers in the user's template are matched to tailored experience
# entries GENERICALLY (no hardcoded company list): by shared years (role date
# ranges are language-independent and near-unique per role) plus shared Latin
# tokens (company names and acronyms usually appear in Latin even in
# Hebrew CVs). A bold non-bullet paragraph containing a 4-digit year is treated
# as a role header; section headings (no years) are left alone.
_YEAR_RE = re.compile(r"(?:19|20)\d{2}")


def _sig(text):
    """(latin_tokens, years) signature of a text, for header<->entry matching."""
    low = (text or "").lower()
    return (set(re.findall(r"[a-z]{3,}", low)), set(_YEAR_RE.findall(low)))


def _match_score(sig_a, sig_b):
    return 2 * len(sig_a[0] & sig_b[0]) + len(sig_a[1] & sig_b[1])


def _strong(ch):
    if _HEB.match(ch):
        return True
    if ch.isascii() and ch.isalpha():
        return False
    return None


def _segments(text):
    """Split into (segment, is_rtl) runs; digits neutral, with the 3PL merge and
    the LTR->RTL boundary-space shift."""
    default = next((s for s in (_strong(c) for c in text) if s is not None), True)
    out, buf, cur = [], "", None
    for ch in text:
        s = _strong(ch)
        kind = s if s is not None else (cur if cur is not None else default)
        if cur is None or kind == cur:
            buf, cur = buf + ch, kind
        else:
            out.append([buf, cur])
            buf, cur = ch, kind
    if buf:
        out.append([buf, cur])

    # merge a Hebrew run's trailing digits into a directly-following Latin run
    # ("אתר 3" + "PL" -> "אתר " + "3PL"), so the number does not reverse.
    for i in range(len(out) - 1):
        seg, rtl = out[i]
        nseg, nrtl = out[i + 1]
        if rtl is True and nrtl is False:
            m = _TRAIL_DIGITS.search(seg)
            if m and not seg.endswith(" "):
                out[i][0] = seg[:m.start()]
                out[i + 1][0] = m.group(0) + nseg

    # move a boundary space from the end of an LTR run to the start of the next
    # RTL run (Word trims trailing whitespace at a direction boundary).
    for i in range(len(out) - 1):
        seg, rtl = out[i]
        nseg, nrtl = out[i + 1]
        if rtl is False and nrtl is True and seg.endswith(" "):
            st = seg.rstrip(" ")
            out[i][0] = st
            out[i + 1][0] = seg[len(st):] + nseg
    return [(s, r) for s, r in out if s]


def _run_size(p):
    for r in p.runs:
        rPr = r._r.find(qn("w:rPr"))
        if rPr is not None and rPr.find(qn("w:sz")) is not None:
            return rPr.find(qn("w:sz")).get(qn("w:val"))
    return None


def _run_bold(p):
    for r in p.runs:
        rPr = r._r.find(qn("w:rPr"))
        if rPr is not None and rPr.find(qn("w:b")) is not None:
            return True
    return False


def refill(p, text):
    """Replace a paragraph's run text with `text`, preserving the paragraph's
    pPr (alignment/style/spacing) and copying size/bold from its current runs."""
    sz = _run_size(p)
    bold = _run_bold(p)
    for r in list(p.runs):
        r._r.getparent().remove(r._r)
    for seg, is_rtl in _segments(text):
        run = p.add_run(seg)
        rPr = run._r.get_or_add_rPr()
        rPr.append(rPr.makeelement(qn("w:rFonts"), {
            qn("w:ascii"): FONT, qn("w:hAnsi"): FONT, qn("w:cs"): FONT}))
        if bold:
            rPr.append(rPr.makeelement(qn("w:b"), {}))
            rPr.append(rPr.makeelement(qn("w:bCs"), {}))
        if sz:
            rPr.append(rPr.makeelement(qn("w:sz"), {qn("w:val"): sz}))
            rPr.append(rPr.makeelement(qn("w:szCs"), {qn("w:val"): sz}))
        if is_rtl:
            rPr.append(rPr.makeelement(qn("w:rtl"), {}))
    return p


def _is_bullet(p):
    # Word-authored bullets carry inline w:numPr; style-based bullets (e.g. a
    # template built programmatically, or some Word styles) only carry a List*
    # paragraph style. Accept both.
    pPr = p._p.find(qn("w:pPr"))
    if pPr is not None and pPr.find(qn("w:numPr")) is not None:
        return True
    try:
        return p.style.name.startswith("List")
    except Exception:
        return False


def _clone_after(p):
    """Deep-copy paragraph p and insert the copy right after it. Returns the new
    Paragraph-like handle (an lxml element wrapper via python-docx)."""
    new_el = copy.deepcopy(p._p)
    p._p.addnext(new_el)
    from docx.text.paragraph import Paragraph
    return Paragraph(new_el, p._parent)


def render_cv_from_template(cv, out_path, template_path=None):
    """Fill the template: replace the professional summary and each role's bullets
    with the tailored text in `cv`. Headers, dates, education, skills, contact and
    name stay verbatim (fixed facts). `cv` is the pipeline's cv spec."""
    if not template_path:
        raise ValueError("hebrew template_path required — set "
                         "'hebrew_cv_template' in config.json")
    doc = Document(template_path)
    allp = doc.paragraphs

    # --- pass 1: locate targets (no mutation, so indices stay valid) ---
    first_bullet_idx = next((i for i, p in enumerate(allp) if _is_bullet(p)), len(allp))
    # summary = the longest Normal paragraph before the first role, excluding the
    # contact line (has @). It's ~350 chars vs <90 for everything else, so length
    # is the unambiguous signal (the paragraph carries a bold run, so a non-bold
    # filter wrongly skips it).
    cands = [p for p in allp[:first_bullet_idx]
             if p.style.name == "Normal" and "@" not in p.text and p.text.strip()]
    summary_p = max(cands, key=lambda p: len(p.text), default=None)

    # Tailored entries with their signatures (all string fields except bullets)
    entries = []
    for e in cv.get("experience", []):
        txt = " ".join(str(v) for k2, v in e.items()
                       if k2 != "bullets" and isinstance(v, str))
        entries.append((_sig(txt), e.get("bullets", []) or []))

    groups = []  # (key, [bullet paragraph objects])
    i = 0
    while i < len(allp):
        p = allp[i]
        is_header = (_run_bold(p) and not _is_bullet(p)
                     and _YEAR_RE.search(p.text or ""))
        if is_header:
            bl, j = [], i + 1
            while j < len(allp) and _is_bullet(allp[j]):
                bl.append(allp[j])
                j += 1
            groups.append((_sig(p.text), bl))
            i = j
        else:
            i += 1

    # --- pass 2: mutate via stable paragraph-object references ---
    if summary_p is not None and cv.get("summary"):
        refill(summary_p, cv["summary"])

    used = set()
    for hsig, bl in groups:
        # best unused tailored entry for this header; require a real signal
        best, best_score = None, 0
        for idx2, (esig, ebullets) in enumerate(entries):
            if idx2 in used:
                continue
            sc = _match_score(hsig, esig)
            if sc > best_score:
                best, best_score = idx2, sc
        if best is None or best_score < 1:
            continue                     # unmatched header keeps template bullets
        used.add(best)
        new = entries[best][1]
        if not bl or not new:
            continue
        for idx in range(min(len(bl), len(new))):
            refill(bl[idx], new[idx])
        for extra in new[len(bl):]:          # more tailored bullets: clone last
            nb = _clone_after(bl[-1])
            refill(nb, extra)
            bl.append(nb)
        for leftover in bl[len(new):]:       # fewer: delete surplus template bullets
            leftover._p.getparent().remove(leftover._p)

    doc.save(out_path)
    return out_path


def render_cover_from_template(cover, out_path, template_path=None):
    """Build the Hebrew cover-email + interview-prep sheet by reusing the
    template's paragraph structures (which render correctly in Word). Clone a
    heading paragraph and a body paragraph as prototypes, clear the body, then
    append cloned+refilled paragraphs for the cover content."""
    if not template_path:
        raise ValueError("hebrew template_path required — set "
                         "'hebrew_cv_template' in config.json")
    doc = Document(template_path)
    paras = doc.paragraphs
    head_proto = next((p for p in paras if p.style.name.startswith("Heading")), paras[0])
    body_proto = next((p for p in paras
                       if p.style.name == "Normal" and not _run_bold(p)
                       and len(p.text.strip()) > 30 and "@" not in p.text), paras[3])
    head_el = copy.deepcopy(head_proto._p)
    body_el = copy.deepcopy(body_proto._p)

    body = doc.element.body
    for p in list(doc.paragraphs):
        p._p.getparent().remove(p._p)

    from docx.text.paragraph import Paragraph

    def add(proto_el, text):
        el = copy.deepcopy(proto_el)
        sectPr = body.find(qn("w:sectPr"))
        if sectPr is not None:
            sectPr.addprevious(el)
        else:
            body.append(el)
        p = Paragraph(el, doc.paragraphs[0]._parent if doc.paragraphs else None)
        refill(p, text)
        return p

    meta = cover.get("meta", {})
    if meta:
        add(head_el, cover.get("headingCover", "פתיח מכתב מועמדות"))
    for k, v in meta.items():
        add(body_el, f"{k}: {v}")
    if cover.get("coverEmail"):
        add(body_el, cover["coverEmail"])
    if cover.get("gaps"):
        add(head_el, cover.get("headingPrep", "הכנה לראיון"))
        for g in cover["gaps"]:
            if g.get("gap"):
                add(body_el, "• " + g["gap"])
            if g.get("framing"):
                add(body_el, f"{g.get('framingLabel', 'מסגור')}: {g['framing']}")
            if g.get("angle"):
                add(body_el, f"{g.get('angleLabel', 'זווית')}: {g['angle']}")

    doc.save(out_path)
    return out_path
